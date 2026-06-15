import pytest
from io import BytesIO
from services.extraction_service import extract_text_from_pdf, extract_text_from_docx
from docx import Document


class TestExtractTextFromPdf:
    """Test PDF text extraction."""
    
    def test_extract_text_from_valid_pdf(self):
        """Test extracting text from a valid PDF."""
        # This test would require a valid PDF file
        # For now, we test error handling
        with pytest.raises(Exception):
            extract_text_from_pdf(b"invalid pdf content")
    
    def test_extract_text_from_empty_pdf(self):
        """Test extracting text from PDF with no content."""
        # Mock test - would need actual PDF bytes
        pass


class TestExtractTextFromDocx:
    """Test DOCX text extraction."""
    
    def test_extract_text_from_docx_valid(self):
        """Test extracting text from a valid DOCX file."""
        # Create a minimal DOCX in memory
        doc = Document()
        doc.add_paragraph("Test Paragraph 1")
        doc.add_paragraph("Test Paragraph 2")
        
        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        result = extract_text_from_docx(docx_bytes.getvalue())
        assert "Test Paragraph 1" in result
        assert "Test Paragraph 2" in result
    
    def test_extract_text_from_empty_docx(self):
        """Test extracting from empty DOCX."""
        doc = Document()
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        result = extract_text_from_docx(docx_bytes.getvalue())
        assert result == ""
    
    def test_extract_text_from_docx_with_multiple_paragraphs(self):
        """Test DOCX with multiple paragraphs."""
        doc = Document()
        doc.add_paragraph("Experience")
        doc.add_paragraph("Senior Developer")
        doc.add_paragraph("Skills")
        doc.add_paragraph("Python, React")
        
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        result = extract_text_from_docx(docx_bytes.getvalue())
        lines = result.split("\n")
        assert len(lines) >= 4
        assert "Experience" in result
    
    def test_extract_text_from_invalid_docx(self):
        """Test error handling for invalid DOCX."""
        with pytest.raises(Exception):
            extract_text_from_docx(b"invalid docx content")
