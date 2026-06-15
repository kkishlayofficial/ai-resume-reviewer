"""
docx_service.py
Generates an editable DOCX or a clean resume PDF from a structured Resume model.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from schemas.resume import Resume


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _has_content(value) -> bool:
    """Return True if the value is non-empty (non-blank string or non-empty list)."""
    if isinstance(value, list):
        return any(_has_content(v) for v in value)
    if isinstance(value, str):
        return bool(value.strip())
    return bool(value)


# ─── DOCX generation ─────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def generate_docx(resume: Resume, output_path: str) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Contact header ────────────────────────────────────────────────────────
    contact = resume.contact
    if _has_content(contact.name):
        name_para = doc.add_heading(contact.name, level=0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in name_para.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    contact_parts = [
        p for p in [
            contact.email,
            contact.phone,
            contact.location,
            contact.linkedin,
            contact.github,
        ] if _has_content(p)
    ]
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9)

    # ── Summary ───────────────────────────────────────────────────────────────
    if _has_content(resume.summary):
        _add_heading(doc, "Professional Summary", level=1)
        doc.add_paragraph(resume.summary)

    # ── Experience ────────────────────────────────────────────────────────────
    if _has_content(resume.experience):
        _add_heading(doc, "Work Experience", level=1)
        for exp in resume.experience:
            if not _has_content(exp.company) and not _has_content(exp.role):
                continue
            title_parts = [p for p in [exp.role, exp.company] if _has_content(p)]
            role_para = doc.add_paragraph()
            run = role_para.add_run(" | ".join(title_parts))
            run.bold = True
            run.font.size = Pt(11)
            if _has_content(exp.duration):
                dur_run = role_para.add_run(f"  {exp.duration}")
                dur_run.font.size = Pt(10)
                dur_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            for bullet in exp.bullets:
                if _has_content(bullet):
                    _add_bullet(doc, bullet.lstrip("•- "))

    # ── Education ─────────────────────────────────────────────────────────────
    if _has_content(resume.education):
        _add_heading(doc, "Education", level=1)
        for edu in resume.education:
            if not _has_content(edu.institution):
                continue
            # Line 1: Institution name (bold) + duration (muted, same line)
            inst_para = doc.add_paragraph()
            inst_run = inst_para.add_run(edu.institution)
            inst_run.bold = True
            inst_run.font.size = Pt(11)
            if _has_content(edu.duration):
                dur_run = inst_para.add_run(f"  |  {edu.duration}")
                dur_run.font.size = Pt(10)
                dur_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            # Line 2: Degree
            if _has_content(edu.degree):
                deg_para = doc.add_paragraph(edu.degree)
                for run in deg_para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            for detail in edu.details:
                if _has_content(detail):
                    _add_bullet(doc, detail)

    # ── Skills ────────────────────────────────────────────────────────────────
    if _has_content(resume.skills):
        _add_heading(doc, "Skills", level=1)
        doc.add_paragraph(", ".join(s for s in resume.skills if _has_content(s)))

    # ── Projects ─────────────────────────────────────────────────────────────
    if _has_content(resume.projects):
        _add_heading(doc, "Projects", level=1)
        for proj in resume.projects:
            if not _has_content(proj.name):
                continue
            proj_para = doc.add_paragraph()
            run = proj_para.add_run(proj.name)
            run.bold = True
            run.font.size = Pt(11)
            if _has_content(proj.description):
                doc.add_paragraph(proj.description)
            for bullet in proj.bullets:
                if _has_content(bullet):
                    _add_bullet(doc, bullet.lstrip("•- "))
            if _has_content(proj.technologies):
                tech_para = doc.add_paragraph()
                t_run = tech_para.add_run("Technologies: ")
                t_run.bold = True
                t_run.font.size = Pt(10)
                tech_para.add_run(", ".join(t for t in proj.technologies if _has_content(t)))

    # ── Certifications ────────────────────────────────────────────────────────
    if _has_content(resume.certifications):
        _add_heading(doc, "Certifications", level=1)
        for cert in resume.certifications:
            if _has_content(cert):
                _add_bullet(doc, cert)

    doc.save(output_path)


# ─── Clean resume PDF generation ─────────────────────────────────────────────

_PDF_PRIMARY = colors.HexColor("#2563EB")
_PDF_TEXT = colors.HexColor("#111827")
_PDF_MUTED = colors.HexColor("#6B7280")
_PDF_BORDER = colors.HexColor("#E5E7EB")

_pdf_styles = getSampleStyleSheet()

_name_style = ParagraphStyle(
    "ResumeName",
    parent=_pdf_styles["Title"],
    fontSize=22,
    textColor=_PDF_PRIMARY,
    spaceAfter=4,
    alignment=1,  # CENTER
)

_contact_style = ParagraphStyle(
    "ResumeContact",
    parent=_pdf_styles["Normal"],
    fontSize=9,
    textColor=_PDF_MUTED,
    spaceAfter=12,
    alignment=1,
)

_section_heading_style = ParagraphStyle(
    "ResumeSectionHeading",
    parent=_pdf_styles["Heading2"],
    fontSize=12,
    textColor=_PDF_PRIMARY,
    spaceBefore=14,
    spaceAfter=6,
    fontName="Helvetica-Bold",
)

_body_style = ParagraphStyle(
    "ResumeBody",
    parent=_pdf_styles["Normal"],
    fontSize=10,
    textColor=_PDF_TEXT,
    leading=16,
    spaceAfter=4,
)

_bold_style = ParagraphStyle(
    "ResumeBold",
    parent=_body_style,
    fontName="Helvetica-Bold",
)

_bullet_style = ParagraphStyle(
    "ResumeBullet",
    parent=_body_style,
    leftIndent=14,
    bulletIndent=4,
    spaceAfter=2,
)

_muted_style = ParagraphStyle(
    "ResumeMuted",
    parent=_body_style,
    fontSize=10,
    textColor=_PDF_MUTED,
    spaceAfter=2,
)


def _sanitize_pdf(text: str) -> str:
    """Escape characters that ReportLab XML parser requires escaping."""
    return (
        text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def generate_resume_pdf(resume: Resume, output_path: str) -> None:
    doc = SimpleDocTemplate(
        output_path,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    story = []

    contact = resume.contact

    # ── Name ──────────────────────────────────────────────────────────────────
    if _has_content(contact.name):
        story.append(Paragraph(_sanitize_pdf(contact.name), _name_style))

    # ── Contact line ──────────────────────────────────────────────────────────
    contact_parts = [
        p for p in [
            contact.email,
            contact.phone,
            contact.location,
            contact.linkedin,
            contact.github,
        ] if _has_content(p)
    ]
    if contact_parts:
        story.append(Paragraph(_sanitize_pdf(" | ".join(contact_parts)), _contact_style))

    def _section(title: str):
        story.append(HRFlowable(width="100%", thickness=0.5, color=_PDF_BORDER, spaceAfter=6))
        story.append(Paragraph(_sanitize_pdf(title), _section_heading_style))

    # ── Summary ───────────────────────────────────────────────────────────────
    if _has_content(resume.summary):
        _section("Professional Summary")
        story.append(Paragraph(_sanitize_pdf(resume.summary), _body_style))

    # ── Experience ────────────────────────────────────────────────────────────
    if _has_content(resume.experience):
        _section("Work Experience")
        for exp in resume.experience:
            if not _has_content(exp.company) and not _has_content(exp.role):
                continue
            title_parts = [p for p in [exp.role, exp.company] if _has_content(p)]
            line = " | ".join(title_parts)
            if _has_content(exp.duration):
                line += f"  |  {exp.duration}"
            story.append(Paragraph(_sanitize_pdf(line), _bold_style))
            for bullet in exp.bullets:
                if _has_content(bullet):
                    story.append(Paragraph(f"• {_sanitize_pdf(bullet.lstrip('•- '))}", _bullet_style))
            story.append(Spacer(1, 4))

    # ── Education ─────────────────────────────────────────────────────────────
    if _has_content(resume.education):
        _section("Education")
        for edu in resume.education:
            if not _has_content(edu.institution):
                continue
            # Line 1: Institution name (bold) + duration (muted, same line)
            inst_line = _sanitize_pdf(edu.institution)
            if _has_content(edu.duration):
                inst_line += f"  |  {_sanitize_pdf(edu.duration)}"
            story.append(Paragraph(inst_line, _bold_style))
            # Line 2: Degree
            if _has_content(edu.degree):
                story.append(Paragraph(_sanitize_pdf(edu.degree), _muted_style))
            for detail in edu.details:
                if _has_content(detail):
                    story.append(Paragraph(f"• {_sanitize_pdf(detail)}", _bullet_style))
            story.append(Spacer(1, 4))

    # ── Skills ────────────────────────────────────────────────────────────────
    if _has_content(resume.skills):
        _section("Skills")
        story.append(Paragraph(
            _sanitize_pdf(", ".join(s for s in resume.skills if _has_content(s))),
            _body_style,
        ))

    # ── Projects ─────────────────────────────────────────────────────────────
    if _has_content(resume.projects):
        _section("Projects")
        for proj in resume.projects:
            if not _has_content(proj.name):
                continue
            story.append(Paragraph(_sanitize_pdf(proj.name), _bold_style))
            if _has_content(proj.description):
                story.append(Paragraph(_sanitize_pdf(proj.description), _body_style))
            for bullet in proj.bullets:
                if _has_content(bullet):
                    story.append(Paragraph(f"• {_sanitize_pdf(bullet.lstrip('•- '))}", _bullet_style))
            if _has_content(proj.technologies):
                tech_line = "<b>Technologies:</b> " + _sanitize_pdf(
                    ", ".join(t for t in proj.technologies if _has_content(t))
                )
                story.append(Paragraph(tech_line, _body_style))
            story.append(Spacer(1, 4))

    # ── Certifications ────────────────────────────────────────────────────────
    if _has_content(resume.certifications):
        _section("Certifications")
        for cert in resume.certifications:
            if _has_content(cert):
                story.append(Paragraph(f"• {_sanitize_pdf(cert)}", _bullet_style))

    doc.build(story)
